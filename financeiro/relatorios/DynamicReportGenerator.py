from docx import Document
import os
import subprocess

class DynamicReportGenerator:
    def __init__(self, json_data, template_path, output_path):
        """
        Inicializa o gerador de relatórios.
        
        :param json_data: Dados JSON para preenchimento do relatório.
        :param template_path: Caminho para o modelo do documento.
        :param output_path: Caminho para salvar o relatório preenchido.
        """
        self.json_data = json_data
        self.template_path = template_path
        self.output_path = output_path
        self.doc = Document(template_path)

    def replace_text_placeholders(self):
        """
        Substitui os campos de texto simples no documento.
        """
        for paragraph in self.doc.paragraphs:
            for key, value in self.json_data.items():
                if key.startswith("$") and isinstance(value, str):
                    paragraph.text = paragraph.text.replace(key, value)

    def fill_tables(self):
        """
        Preenche tabelas no documento com base no JSON.
        """
        for table in self.doc.tables:
            table_name = None

            # Detectar o nome da tabela com base no primeiro marcador
            for row in table.rows:
                for cell in row.cells:
                    for key in self.json_data.get("tabelas", {}).keys():
                        if any(marker in cell.text for marker in self.json_data["tabelas"][key][0].keys()):
                            table_name = key
                            break
                    if table_name:
                        break
                if table_name:
                    break

            # Preencher a tabela correspondente
            if table_name:
                data_rows = self.json_data["tabelas"][table_name]
                for data_row in data_rows:
                    new_row = table.add_row().cells
                    col_index = 0
                    for key, value in data_row.items():
                        if col_index < len(new_row):
                            new_row[col_index].text = (
                                f"R$ {value:,.2f}".replace(",", "#").replace(".", ",").replace("#", ".")
                                if isinstance(value, float)
                                else str(value)
                            )
                            col_index += 1

                # Remover a linha de exemplo com marcadores
                for row in table.rows:
                    for cell in row.cells:
                        if any(marker in cell.text for marker in data_rows[0].keys()):
                            cell.text = ""

            # Substituir variáveis em tabelas menores
            else:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in self.json_data.items():
                            if key.startswith("$"):# and isinstance(value, str):
                                formatted_value = (
                                    f"R$ {value:,.2f}".replace(",", "#").replace(".", ",").replace("#", ".")
                                    if isinstance(value, (int, float))
                                    else str(value)
                                )
                                cell.text = cell.text.replace(key, formatted_value)
                                

    def generate_report(self):
        """
        Gera o relatório preenchido.
        """
        self.replace_text_placeholders()
        self.fill_tables()
        self.doc.save(self.output_path)
        print(f"Relatório gerado com sucesso: {self.output_path}")
        # Converter para PDF
        self.convert_to_pdf()


    # Converter ODT para PDF usando LibreOffice
    def convert_to_pdf(self):
        try:
            subprocess.run([
                #"soffice" #Usar em Linux
                r"C:\Program Files\LibreOffice\program\soffice.exe",  #Usar em Windows
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(self.output_path),#Só pega o diretório nome é o mesmo só que com .pdf
                self.output_path
            ], check=True)
            print(f"PDF gerado com sucesso!")
        except FileNotFoundError:
            print("Erro: O LibreOffice não foi encontrado. Certifique-se de que está instalado e configurado no PATH do sistema.")
        except subprocess.CalledProcessError as e:
            print(f"Erro ao converter para PDF: {e}")



'''
# Caminho dinâmico baseado no diretório atual
current_dir = os.path.dirname(os.path.abspath(__file__))
template_path = os.path.join(current_dir, "modelo_dre_mensal.docx")
output_path = os.path.join(current_dir, "saida\dre_mensal_preenchido.docx")

# JSON de entrada
json_data = {
    "$data": "16/01/2025",
    "tabelas": {
        "accounts": [
            {"$conta": "Mercado Pago - PF", "$saldo": 1000.0, "$status": "Ativo"},
            {"$conta": "XP - PF", "$saldo": 500.0, "$status": "Bloqueado"}
        ],
        "lancamentos": [
            {"$lanc_dt": "15/01/2025", "$lanc_valor": 300.62, "$lanc_origem": "Origem 3", "$lanc_descricao": "Descrição do lançamento 1", "$pa" : "", "$pt": ""},
            {"$lanc_dt": "14/01/2025", "$lanc_valor": 563.67, "$lanc_origem": "Origem 4", "$lanc_descricao": "Descrição do lançamento 2", "$pa" : "", "$pt": ""}
        ]
    },
    "$total_entradas": 1200.0, 
    "$total_despesas": 1000.0, 
    "$total_investimento": 50.0, 
    "$total_saldo": 150.0
}

# Gerar relatório
report_generator = DynamicReportGenerator(json_data, template_path, output_path)
report_generator.generate_report()

'''